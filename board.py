
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Colour palette ────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1A, 0x35, 0x5E)   # dark navy – headings
C_BLUE   = RGBColor(0x26, 0x6D, 0xC4)   # mid blue – sub-headings
C_LTBLUE = RGBColor(0xD9, 0xE8, 0xF8)   # light blue – header rows
C_RED    = RGBColor(0xC0, 0x39, 0x2B)   # high priority
C_AMBER  = RGBColor(0xE6, 0x7E, 0x22)   # medium priority
C_GREEN  = RGBColor(0x1E, 0x8B, 0x4C)   # low priority
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_LGRAY  = RGBColor(0xF4, 0xF6, 0xF9)   # alt row fill
C_DKGRY  = RGBColor(0x44, 0x44, 0x55)   # body text

# ── Helper utilities ──────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  str(rgb))
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='AAAAAA', sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_shading(para, fill_hex):
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)

def add_para_border_bottom(para, color='2F6DC4', sz=6):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_run_font(run, name='Calibri'):
    run.font.name = name
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    name)
    rFonts.set(qn('w:hAnsi'),    name)
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

# ── Title page ────────────────────────────────────────────────────────────────
def add_cover():
    # Spacer
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('Mansa Dashboard')
    run.font.size  = Pt(34)
    run.font.bold  = True
    run.font.color.rgb = C_NAVY
    set_run_font(run)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run('Analysis & Improvement Plan')
    r2.font.size  = Pt(22)
    r2.font.bold  = False
    r2.font.color.rgb = C_BLUE
    set_run_font(r2)

    # Divider line via paragraph bottom border
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_before = Pt(4)
    div.paragraph_format.space_after  = Pt(4)
    add_para_border_bottom(div, color='2F6DC4', sz=12)

    # Sub-line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run('Architecture Review  ·  Tabbly Dependency Reduction  ·  Performance & UX Improvements')
    r3.font.size  = Pt(11)
    r3.font.color.rgb = C_DKGRY
    r3.font.italic = True
    set_run_font(r3)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(0)
    r4 = p4.add_run('Prepared for: Mansa Infotech  ·  April 2026')
    r4.font.size  = Pt(10)
    r4.font.color.rgb = C_DKGRY
    set_run_font(r4)

    # Page break
    doc.add_page_break()

# ── Section heading ───────────────────────────────────────────────────────────
def add_section_heading(emoji, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'{emoji}  {number}. {title}')
    run.font.size  = Pt(15)
    run.font.bold  = True
    run.font.color.rgb = C_NAVY
    set_run_font(run)
    add_para_border_bottom(p, color='1A355E', sz=8)
    return p

def add_sub_heading(title, emoji=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    label = f'{emoji}  {title}' if emoji else title
    run = p.add_run(label)
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = C_BLUE
    set_run_font(run)

def add_body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = C_DKGRY
    set_run_font(run)
    return p

def add_bullet(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' ')
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.color.rgb = C_NAVY
        set_run_font(rb)
    r = p.add_run(text)
    r.font.size  = Pt(10.5)
    r.font.color.rgb = C_DKGRY
    set_run_font(r)

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3)
    set_para_shading(p, 'EFF3FA')
    run = p.add_run(text)
    run.font.name   = 'Courier New'
    run.font.size   = Pt(9)
    run.font.color.rgb = RGBColor(0x19, 0x32, 0x66)

# ── Priority badge cell helper ────────────────────────────────────────────────
PRIORITY_COLORS = {
    '🔴 High':   (RGBColor(0xFD, 0xED, 0xED), C_RED),
    '🟡 Medium': (RGBColor(0xFF, 0xF6, 0xE5), C_AMBER),
    '🟢 Low':    (RGBColor(0xE8, 0xF6, 0xEE), C_GREEN),
}

def styled_table(headers, rows, col_widths):
    """Create a polished table with alternating rows."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for i, w in enumerate(col_widths):
        for cell in tbl.columns[i].cells:
            cell.width = Inches(w)

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_NAVY)
        set_cell_border(cell, color='FFFFFF', sz=4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = C_WHITE
        set_run_font(run)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        bg  = C_LGRAY if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_border(cell, color='CCCCCC', sz=3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Priority column special styling
            if val in PRIORITY_COLORS:
                bg_c, fg_c = PRIORITY_COLORS[val]
                set_cell_bg(cell, bg_c)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                run = p.add_run(val)
                run.font.bold  = True
                run.font.size  = Pt(9)
                run.font.color.rgb = fg_c
                set_run_font(run)
            else:
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                run = p.add_run(val)
                run.font.size  = Pt(9.5)
                run.font.color.rgb = C_DKGRY
                set_run_font(run)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

add_cover()

# ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
run = p.add_run('Executive Summary')
run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = C_NAVY
set_run_font(run)
add_para_border_bottom(p, color='1A355E', sz=6)

add_body(
    'This document presents a comprehensive analysis of the Mansa Dashboard codebase — a FastAPI + React '
    'application used to manage AI calling agents, trigger outbound calls, track call logs, and automate '
    'Cal.com meeting bookings via Tabbly.io. The review covers the full backend and frontend source code '
    'and identifies critical issues around security, external API dependency, performance, reliability, '
    'and user experience. Each section includes actionable recommendations with priority ratings.'
)

# ── SECTION 1 ─────────────────────────────────────────────────────────────────
add_section_heading('🔍', 1, 'Current System Analysis')

add_sub_heading('Architecture Overview', '🏗️')
add_body(
    'The system is a monolith: a single FastAPI server simultaneously serves the React frontend (built '
    'static files) and all backend API routes under /api/*. Supabase handles authentication and stores '
    'two internal tables — agent_mappings and meeting_logs. All other data (agents, call logs, campaigns) '
    'is fetched live from Tabbly.io on every request.'
)

add_sub_heading('Data Flow')
add_bullet('User logs in via Supabase Auth (JWT issued)')
add_bullet('Frontend sends JWT in Authorization header to FastAPI')
add_bullet('FastAPI validates JWT against Supabase, extracts user_id')
add_bullet('Routers call Tabbly.io APIs synchronously and return responses')
add_bullet('Post-call processing uses background tasks with time.sleep(120)')

add_sub_heading('Key External Dependencies')
styled_table(
    ['Service', 'Used For', 'Criticality'],
    [
        ['Tabbly.io', 'Agent CRUD, call triggering, call logs, campaigns', '🔴 Critical — no fallback'],
        ['Supabase', 'Auth, agent_mappings, meeting_logs', '🔴 Critical'],
        ['Cal.com', 'Availability fetching, meeting booking', '🟡 Important'],
        ['Google Gemini', 'Transcript parsing for email/date extraction', '🟡 Important'],
    ],
    [2.2, 4.0, 2.2]
)

add_sub_heading('Identified Bottlenecks')
add_bullet('Every dashboard load triggers 4 sequential API calls: fetchAgents → fetchCallLogs → fetchStats → getMeetingLogs')
add_bullet('Call Logs page polls Tabbly every 5 seconds, each poll fetching logs for every agent')
add_bullet('Stats endpoint fetches 100 logs per agent from Tabbly live on every request — O(n×agents) external calls')
add_bullet('post_call_service.py blocks for 120 seconds using time.sleep() inside a background task (not a queue/worker)')
add_bullet('Meeting-logs endpoint makes two sequential Tabbly calls: one for logs, one for agent names')

# ── SECTION 2 ─────────────────────────────────────────────────────────────────
add_section_heading('⚙️', 2, 'Reducing Dependency on External APIs')

add_sub_heading('Current Problem')
add_body(
    'The dashboard has zero local storage of call data. Every piece of call history, transcript, and '
    'recording URL lives exclusively in Tabbly. If Tabbly is slow, rate-limits the account, or goes '
    'offline, the entire dashboard breaks. There is no cache, no fallback, and no offline mode.'
)

add_sub_heading('Recommended: Local call_logs Table in Supabase')
add_body('Add the following table to capture all call events internally:')
add_code('CREATE TABLE call_logs (')
add_code('  id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),')
add_code('  user_id         UUID REFERENCES auth.users(id),')
add_code('  agent_id        TEXT NOT NULL,')
add_code('  call_id         TEXT UNIQUE NOT NULL,   -- Tabbly participant_identity')
add_code('  phone_number    TEXT,')
add_code('  status          TEXT,                   -- Completed / Not Answered / Processing')
add_code('  transcript      TEXT,')
add_code('  recording_url   TEXT,')
add_code('  json_output     JSONB,')
add_code('  duration_secs   INT DEFAULT 0,')
add_code('  called_at       TIMESTAMPTZ,')
add_code('  created_at      TIMESTAMPTZ DEFAULT now()')
add_code(');')
add_code('CREATE INDEX idx_cl_user_date  ON call_logs(user_id, called_at DESC);')
add_code('CREATE INDEX idx_cl_agent      ON call_logs(agent_id);')
add_code('CREATE INDEX idx_cl_status     ON call_logs(status);')

add_sub_heading('Recommended: Local campaigns Table')
add_code('CREATE TABLE campaigns (')
add_code('  id                 SERIAL PRIMARY KEY,')
add_code('  user_id            UUID REFERENCES auth.users(id),')
add_code('  tabbly_campaign_id INT,')
add_code('  campaign_name      TEXT,')
add_code('  agent_id           TEXT,')
add_code('  status             TEXT DEFAULT \'pending\',')
add_code('  start_time         TIMESTAMPTZ,')
add_code('  end_time           TIMESTAMPTZ,')
add_code('  created_at         TIMESTAMPTZ DEFAULT now()')
add_code(');')

add_sub_heading('Migration Strategy')
add_bullet('Phase 1 (Week 1–2): Write call data to Supabase via the existing /api/webhooks/tabbly endpoint when calls end')
add_bullet('Phase 2 (Week 3–4): Modify GET /api/logs/call-logs to query Supabase first; fall back to Tabbly only for fresh calls')
add_bullet('Phase 3 (Week 5–6): Stats, filtering, and analytics now fully served from local DB — Tabbly used for writes only')
add_bullet('Phase 4 (Ongoing): Run a nightly sync job to backfill any gaps from Tabbly into Supabase')

add_sub_heading('Caching Layer (Redis / In-Memory)')
add_body('Even before full local storage, adding a simple TTL cache reduces Tabbly load by ~80%:')
add_code('# Simple in-memory cache (upgrade to Redis for multi-instance)')
add_code('import time')
add_code('_cache = {}')
add_code('')
add_code('def get_agents_cached(user_id, ttl=60):')
add_code('    key = f"agents:{user_id}"')
add_code('    if key in _cache and time.time() - _cache[key]["ts"] < ttl:')
add_code('        return _cache[key]["data"]')
add_code('    data = tabbly.get_agents()')
add_code('    _cache[key] = {"data": data, "ts": time.time()}')
add_code('    return data')

# ── SECTION 3 ─────────────────────────────────────────────────────────────────
add_section_heading('📊', 3, 'Data Ownership & Analytics')

add_sub_heading('Proposed Analytics Schema')
add_body('Once call data is stored locally, the following SQL queries power a full analytics module:')
add_code('-- Daily call performance per agent')
add_code('SELECT agent_id, DATE(called_at) as day,')
add_code('       COUNT(*) as total_calls,')
add_code('       SUM(CASE WHEN status=\'Completed\' THEN 1 ELSE 0 END) as completed,')
add_code('       AVG(duration_secs) as avg_duration_secs')
add_code('FROM call_logs')
add_code('WHERE user_id = $1')
add_code('GROUP BY agent_id, DATE(called_at)')
add_code('ORDER BY day DESC;')
add_body('')
add_code('-- Meeting conversion funnel')
add_code('SELECT')
add_code('  COUNT(*) FILTER (WHERE status=\'booked\')  AS booked,')
add_code('  COUNT(*) FILTER (WHERE status=\'failed\')  AS failed,')
add_code('  COUNT(*) FILTER (WHERE status=\'skipped\') AS skipped,')
add_code('  COUNT(*) FILTER (WHERE is_interested=true) AS interested')
add_code('FROM meeting_logs')
add_code('WHERE user_id = $1;')

add_sub_heading('Suggested Dashboard Widgets')
styled_table(
    ['Widget', 'Data Source', 'Chart Type'],
    [
        ['Daily Call Volume', 'call_logs grouped by day', 'Bar chart (Recharts)'],
        ['Success Rate Trend', 'completed / total per day', 'Line chart'],
        ['Meeting Conversion Funnel', 'meeting_logs status counts', 'Funnel / Donut chart'],
        ['Top Performing Agents', 'completed calls per agent_id', 'Horizontal bar chart'],
        ['Avg Call Duration', 'AVG(duration_secs) per agent', 'Bar chart'],
        ['Interest vs Booked Rate', 'is_interested vs booked', 'Grouped bar chart'],
    ],
    [2.8, 2.8, 2.8]
)

add_sub_heading('Report Export')
add_bullet('Add a GET /api/analytics/export?format=csv&from=YYYY-MM-DD&to=YYYY-MM-DD endpoint')
add_bullet('Use Python csv module to stream response — no file storage needed')
add_bullet('Frontend: download button on the analytics page triggers this endpoint')

# ── SECTION 4 ─────────────────────────────────────────────────────────────────
add_section_heading('⚡', 4, 'Performance Optimization')

add_sub_heading('Frontend: Parallelize API Calls')
add_body('Currently Dashboard.tsx awaits each API call sequentially. Replace with Promise.all:')
add_code('// ❌ Current — sequential (slow)')
add_code('await fetchAgents();')
add_code('const logsData = await fetchCallLogs(5);')
add_code('const statsData = await fetchStats();')
add_code('const mLogs    = await getMeetingLogs();')
add_body('')
add_code('// ✅ Improved — parallel (fast)')
add_code('const [logsData, statsData, mLogs] = await Promise.all([')
add_code('  fetchCallLogs(5),')
add_code('  fetchStats(),')
add_code('  getMeetingLogs(),')
add_code(']);')
add_body('This single change can reduce perceived dashboard load time by 50–70%.')

add_sub_heading('Frontend: Reduce Polling Frequency')
add_body(
    'CallLogs.tsx polls every 5 seconds. This fires a Tabbly API call every 5 seconds per user. '
    'Replace with Supabase Realtime subscriptions once call_logs are stored locally:'
)
add_code('// Subscribe to new rows in call_logs table')
add_code('supabase')
add_code('  .channel("call_logs")')
add_code('  .on("postgres_changes", { event: "INSERT", schema: "public", table: "call_logs" },')
add_code('    (payload) => setLogs(prev => [payload.new, ...prev]))')
add_code('  .subscribe();')

add_sub_heading('Backend: Add Pagination')
add_body('The /api/logs/call-logs endpoint fetches up to 100 logs per agent with no pagination. Add cursor-based pagination:')
add_code('# Backend')
add_code('@router.get("/call-logs")')
add_code('def get_call_logs(limit: int = 20, offset: int = 0, user_id = Depends(get_user_id)):')
add_code('    logs = supabase.table("call_logs").select("*")')
add_code('                   .eq("user_id", user_id)')
add_code('                   .order("called_at", desc=True)')
add_code('                   .range(offset, offset + limit - 1).execute()')
add_code('    return {"logs": logs.data, "total": logs.count}')

add_sub_heading('Backend: Fix Stats Endpoint')
add_body(
    'Currently GET /api/logs/stats fetches 100 live Tabbly logs per agent on every call. '
    'Once call_logs is local, replace with a single aggregated query:'
)
add_code('SELECT')
add_code('  COUNT(*)                                        AS total_calls,')
add_code('  COUNT(*) FILTER (WHERE status=\'Completed\')   AS total_completed,')
add_code('  COUNT(DISTINCT agent_id)                       AS active_agents')
add_code('FROM call_logs WHERE user_id = $1;')

add_sub_heading('Summary of Performance Improvements')
styled_table(
    ['Improvement', 'Impact', 'Effort'],
    [
        ['Parallelize Dashboard API calls',        'Load time ↓ 50–70%',          'Low (1 hour)'],
        ['Cache Tabbly agent list (60s TTL)',       'Tabbly calls ↓ 80%',          'Low (2 hours)'],
        ['Local call_logs DB + remove live polls', 'Response time ↓ 90%',         'Medium (3–5 days)'],
        ['Supabase Realtime instead of polling',   'Server load ↓ 95%',           'Medium (1–2 days)'],
        ['Pagination (20 rows default)',            'Page weight ↓ 80%',           'Low (2–3 hours)'],
        ['Aggregated SQL stats query',             'Stats load time ↓ 95%',       'Low (2–3 hours)'],
        ['Redis cache for all Tabbly responses',   'P95 latency ↓ ~200ms',        'Medium (2–3 days)'],
    ],
    [3.2, 2.4, 2.0]
)

# ── SECTION 5 ─────────────────────────────────────────────────────────────────
add_section_heading('🛡️', 5, 'Reliability & System Stability')

add_sub_heading('Critical: Replace time.sleep() with a Background Queue')
add_body(
    'post_call_service.py calls time.sleep(120) inside a FastAPI BackgroundTask. This blocks a thread '
    'for 2+ minutes per call, and has no retry tracking, no dead-letter queue, and no visibility. '
    'Replace with ARQ (async Redis queue) or Celery:'
)
add_code('# Install: pip install arq')
add_code('')
add_code('# worker.py')
add_code('async def process_call_task(ctx, call_id: str, agent_id: str, user_id: str):')
add_code('    await asyncio.sleep(120)  # non-blocking async sleep')
add_code('    await post_call_service.process_call_results(call_id, agent_id, user_id)')
add_code('')
add_code('class WorkerSettings:')
add_code('    functions = [process_call_task]')
add_code('    redis_settings = RedisSettings(host="localhost", port=6379)')
add_code('')
add_code('# In router — enqueue instead of background_tasks.add_task:')
add_code('await redis.enqueue_job("process_call_task", call_id, agent_id, user_id)')

add_sub_heading('Add Retry Logic for Tabbly Calls')
add_body('Wrap all Tabbly HTTP requests in a retry decorator:')
add_code('from tenacity import retry, stop_after_attempt, wait_exponential')
add_code('')
add_code('@retry(stop=stop_after_attempt(3),')
add_code('       wait=wait_exponential(multiplier=1, min=2, max=10))')
add_code('def fetch_call_logs(agent_id, limit=50):')
add_code('    ...  # existing Tabbly call')

add_sub_heading('Circuit Breaker Pattern')
add_body(
    'If Tabbly fails 3 times in a row, stop hitting it for 60 seconds and return cached/empty data. '
    'Use the pybreaker library or implement a simple state machine. This prevents cascading failures '
    'from bringing down the entire dashboard.'
)

add_sub_heading('Fix API Keys in Source Code')
add_body(
    'Both tabbly.py and middleware/auth.py have real API keys hardcoded as fallback values. '
    'This is a critical security issue — anyone with repo access has full Tabbly and Supabase access.'
)
add_code('# ❌ Current (DANGEROUS)')
add_code('TABBLY_API_KEY = os.getenv("TABBLY_API_KEY") or "070ef3afd90b3ca7"')
add_code('')
add_code('# ✅ Fixed')
add_code('TABBLY_API_KEY = os.getenv("TABBLY_API_KEY")')
add_code('if not TABBLY_API_KEY:')
add_code('    raise ValueError("TABBLY_API_KEY environment variable is required")')

add_sub_heading('Fix CORS Configuration')
add_code('# ❌ Current (allows all origins)')
add_code('allow_origins=["*"]')
add_code('')
add_code('# ✅ Fixed')
add_code('allow_origins=[os.getenv("FRONTEND_URL", "http://localhost:5173")]')

add_sub_heading('Add Structured Logging & Monitoring')
add_bullet('Replace print() statements throughout services with Python logging module')
add_bullet('Add Sentry for error tracking (pip install sentry-sdk[fastapi])')
add_bullet('Add /api/health endpoint that checks Supabase connectivity and Tabbly reachability')
add_bullet('Use Render/Railway built-in metrics or add Prometheus + Grafana for production monitoring')

add_sub_heading('Fix requirements.txt')
add_body('The current requirements.txt only lists 3 packages. The actual dependencies include:')
add_code('fastapi>=0.111.0')
add_code('uvicorn[standard]>=0.30.0')
add_code('python-dotenv>=1.0.0')
add_code('supabase>=2.4.0')
add_code('requests>=2.31.0')
add_code('google-generativeai>=0.5.0')
add_code('tenacity>=8.2.0          # retry logic')
add_code('arq>=0.25.0              # async background jobs')
add_code('sentry-sdk[fastapi]>=1.40.0')
add_code('urllib3<2')

# ── SECTION 6 ─────────────────────────────────────────────────────────────────
add_section_heading('👨‍💻', 6, 'User Experience (UX) Improvements')

add_sub_heading('Skeleton Loaders Instead of Full-Page Spinners')
add_body(
    'Currently every page shows a centered spinner while all data loads. Skeleton screens show '
    'the page structure immediately, making the app feel significantly faster:'
)
add_code('// Skeleton row component')
add_code('const SkeletonRow = () => (')
add_code('  <tr className="animate-pulse">')
add_code('    <td><div className="h-4 bg-muted/30 rounded w-24 my-2" /></td>')
add_code('    <td><div className="h-4 bg-muted/30 rounded w-32 my-2" /></td>')
add_code('    <td><div className="h-4 bg-muted/30 rounded w-16 my-2" /></td>')
add_code('  </tr>')
add_code(');')

add_sub_heading('Optimistic UI for Call Triggering')
add_body(
    'When a user triggers a call, immediately add a "Processing" row to the call logs table '
    'before the API responds. Remove it only if the API returns an error. This eliminates the '
    'disconnect between "call triggered" and "call appears in logs".'
)

add_sub_heading('Error Boundaries')
add_body('Wrap each page section so a single failed API call does not blank the whole page:')
add_code('// React Error Boundary')
add_code('class ErrorBoundary extends React.Component {')
add_code('  state = { hasError: false };')
add_code('  static getDerivedStateFromError() { return { hasError: true }; }')
add_code('  render() {')
add_code('    if (this.state.hasError)')
add_code('      return <div className="p-6 text-error">Failed to load this section.</div>;')
add_code('    return this.props.children;')
add_code('  }')
add_code('}')

add_sub_heading('Additional UX Recommendations')
add_bullet('Date/time display', 'Show "2 minutes ago" style relative times for recent calls — clearer than raw UTC timestamps')
add_bullet('Search & filter bar', 'Add phone number / agent / status filter to Call Logs — critical once logs grow beyond 50 entries')
add_bullet('Call log date columns', 'Currently displayed in UTC; should show in IST with timezone label (calls are Indian-market focused)')
add_bullet('Campaign status', 'Campaign page has no local state — user cannot see past campaigns without hitting Tabbly. Fix with local campaigns table')
add_bullet('Toast on meeting booked', 'Background jobs should update Supabase Realtime so user gets a toast notification when a meeting is auto-booked')
add_bullet('Mobile responsiveness', 'Tables on CallLogs and MeetingLogs overflow on mobile. Use horizontal scroll containers or card-based layouts for small screens')

# ── SECTION 7 ─────────────────────────────────────────────────────────────────
add_section_heading('🚀', 7, 'Scalability & Future Readiness')

add_sub_heading('Current Scale Limitations')
add_bullet('Single FastAPI process — no horizontal scaling; one server handles all requests and background tasks')
add_bullet('Background tasks run in-process — a restart kills all pending post-call processing jobs')
add_bullet('No rate limiting on API endpoints — a single user can flood Tabbly with unlimited calls')
add_bullet('Tabbly is a hard ceiling — the platform cannot support more agents or calls than Tabbly allows')

add_sub_heading('Recommended Architecture for Scale')
add_body('Separate the system into clear service boundaries:')
add_bullet('API Service', '— FastAPI, stateless, horizontally scalable. Handles auth, CRUD, serves frontend.')
add_bullet('Worker Service', '— ARQ/Celery workers processing post-call jobs. Scale independently.')
add_bullet('Data Layer', '— Supabase (PostgreSQL) as single source of truth. Add read replicas for heavy analytics.')
add_bullet('Cache Layer', '— Redis for Tabbly response caching, session data, rate limiting.')
add_bullet('Tabbly (Vendor)', '— Used only for call triggering and agent CRUD. Not on the read path.')

add_sub_heading('Rate Limiting')
add_code('# Add to main.py using slowapi')
add_code('from slowapi import Limiter')
add_code('from slowapi.util import get_remote_address')
add_code('')
add_code('limiter = Limiter(key_func=get_remote_address)')
add_code('')
add_code('@router.post("/trigger-call")')
add_code('@limiter.limit("10/minute")   # max 10 calls per minute per IP')
add_code('def trigger_call(request: Request, ...):')
add_code('    ...')

add_sub_heading('Environment-based Configuration')
add_body('Introduce proper config management using Pydantic Settings:')
add_code('from pydantic_settings import BaseSettings')
add_code('')
add_code('class Settings(BaseSettings):')
add_code('    tabbly_api_key: str')
add_code('    tabbly_org_id: str')
add_code('    supabase_url: str')
add_code('    supabase_anon_key: str')
add_code('    redis_url: str = "redis://localhost:6379"')
add_code('    frontend_url: str = "http://localhost:5173"')
add_code('    environment: str = "development"')
add_code('')
add_code('    class Config:')
add_code('        env_file = ".env"')

add_sub_heading('Future: Webhook-First Architecture')
add_body(
    'The current dual approach (proactive polling + webhook) creates race conditions. '
    'Long-term, invest in making Tabbly webhooks reliable and remove the proactive polling entirely. '
    'All post-call processing should be event-driven via webhooks → job queue → worker.'
)

# ── SECTION 8: MASTER PRIORITY TABLE ─────────────────────────────────────────
add_section_heading('📌', 8, 'Consolidated Issue Registry & Implementation Roadmap')

add_sub_heading('All Issues — Prioritised')
styled_table(
    ['#', 'Issue', 'Area', 'Priority', 'Est. Effort'],
    [
        ['1',  'API keys hardcoded as fallbacks in tabbly.py & auth.py',                    'Security',       '🔴 High',   '1 hour'],
        ['2',  'CORS allow_origins=["*"] in production',                                    'Security',       '🔴 High',   '30 mins'],
        ['3',  'No local storage of call logs — 100% dependent on Tabbly reads',            'Data Ownership', '🔴 High',   '3–5 days'],
        ['4',  'Dashboard stats queries Tabbly live on every load',                         'Performance',    '🔴 High',   '1–2 days'],
        ['5',  'CallLogs polls Tabbly every 5 seconds per user',                            'Performance',    '🔴 High',   '1–2 days'],
        ['6',  'time.sleep(120) blocks background task thread per call',                    'Reliability',    '🔴 High',   '2–3 days'],
        ['7',  'No retry / circuit-breaker around any Tabbly HTTP call',                    'Reliability',    '🔴 High',   '1 day'],
        ['8',  'requirements.txt missing most actual dependencies',                          'DevOps',         '🔴 High',   '30 mins'],
        ['9',  'Dashboard loads 4 API calls sequentially instead of parallel',              'Performance',    '🟡 Medium', '1 hour'],
        ['10', 'No pagination — loads all logs at once',                                    'Performance',    '🟡 Medium', '3–4 hours'],
        ['11', 'Agent names fetched from Tabbly on every meeting-logs view',                'Performance',    '🟡 Medium', '2 hours'],
        ['12', 'Schema errors silently swallowed with print() in supabase_service.py',      'Reliability',    '🟡 Medium', '2 hours'],
        ['13', 'No structured logging — print() scattered throughout services',             'Reliability',    '🟡 Medium', '1 day'],
        ['14', 'No health check endpoint',                                                  'Reliability',    '🟡 Medium', '2 hours'],
        ['15', 'Campaigns have no local state — lost if Tabbly changes',                    'Data Ownership', '🟡 Medium', '2–3 days'],
        ['16', 'No error boundaries in React — one failure blanks whole page',              'UX',             '🟡 Medium', '2–3 hours'],
        ['17', 'Full-page spinners with no skeleton loaders',                               'UX',             '🟡 Medium', '1 day'],
        ['18', 'No search/filter on Call Logs or Meeting Logs',                             'UX',             '🟡 Medium', '1 day'],
        ['19', 'Times shown in UTC; India-focused app should default to IST',              'UX',             '🟡 Medium', '2 hours'],
        ['20', 'No analytics charts — only raw numbers in stats',                           'Analytics',      '🟡 Medium', '3–5 days'],
        ['21', 'No rate limiting on trigger-call endpoint',                                 'Security',       '🟡 Medium', '2 hours'],
        ['22', 'Call status determined by fragile minutes_ago heuristic',                  'Reliability',    '🟢 Low',    '3 hours'],
        ['23', 'Plus icon component defined inline in Dashboard.tsx',                       'Code Quality',   '🟢 Low',    '30 mins'],
        ['24', 'No mobile-responsive table layouts',                                        'UX',             '🟢 Low',    '1–2 days'],
        ['25', 'No report export (CSV/PDF) for call or meeting logs',                       'Analytics',      '🟢 Low',    '2–3 days'],
    ],
    [0.35, 3.7, 1.5, 1.1, 1.0]
)

add_sub_heading('Suggested Implementation Sprints')
styled_table(
    ['Sprint', 'Duration', 'Focus', 'Key Deliverables'],
    [
        ['Sprint 1 — Security & Quick Wins',   '1 week',   'Security + Performance',    'Remove hardcoded keys, fix CORS, parallelize frontend calls, fix requirements.txt, add pagination'],
        ['Sprint 2 — Data Independence',       '2 weeks',  'Local Storage',             'call_logs table, webhook writes to DB, /logs/call-logs reads from Supabase, remove Tabbly read dependency'],
        ['Sprint 3 — Reliability',             '1 week',   'Queue + Retry',             'ARQ job queue, tenacity retry on all Tabbly calls, circuit breaker, structured logging, /health endpoint'],
        ['Sprint 4 — Analytics & UX',          '2 weeks',  'Analytics + UX polish',     'Analytics charts, Supabase Realtime, skeleton loaders, error boundaries, IST timezone fix, search/filter'],
        ['Sprint 5 — Scale Prep',              '1 week',   'Scalability',               'Redis caching, rate limiting, Pydantic settings, monitoring/alerting setup'],
    ],
    [2.0, 1.0, 1.8, 3.8]
)

# ── FOOTER ────────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para_border_bottom(p, color='1A355E', sz=6)
r = p.add_run('End of Report  ·  Mansa Dashboard Analysis  ·  April 2026')
r.font.size = Pt(9)
r.font.color.rgb = C_DKGRY
r.font.italic = True
set_run_font(r)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = 'Mansa_Dashboard_Analysis_Report.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
